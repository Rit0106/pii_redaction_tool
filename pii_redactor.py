#!/usr/bin/env python3
"""
PII Redaction Tool
==================
Main script that orchestrates PII detection and redaction for .docx documents.

Approach: Hybrid detection using:
  1. Regex patterns for structured PII (emails, phones, SSNs, credit cards, IPs, DOBs)
  2. spaCy NER for unstructured PII (names, organizations, locations)
  3. Context-aware address detection (street keywords + PIN codes + state names)

Usage:
    python pii_redactor.py <input.docx> [--output <output.docx>] [--report]

Author: PII Redaction Tool
"""

import argparse
import copy
import json
import os
import re
import sys
import time
from typing import List, Dict, Tuple, Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from detectors import (
    PIIEntity, RegexDetector, NERDetector, AddressDetector,
    merge_overlapping_entities
)
from replacer import ConsistentReplacer


class PIIRedactor:
    """
    Main redaction engine that coordinates detection and replacement.

    Architecture:
    ┌──────────────┐     ┌──────────────┐     ┌──────────────┐
    │  .docx Input │────>│  Detection   │────>│  Replacement │────> Redacted .docx
    └──────────────┘     │  Pipeline    │     │  Engine      │
                         │              │     └──────────────┘
                         │ • Regex      │
                         │ • NER        │
                         │ • Address    │
                         │ • Context    │
                         └──────────────┘
    """

    def __init__(self, spacy_model: str = "en_core_web_sm"):
        """
        Initialize the redactor with all detection and replacement engines.

        Args:
            spacy_model: Name of the spaCy model for NER (default: en_core_web_sm)
        """
        print("[*] Initializing PII Redaction Engine...")

        # Initialize detectors
        self.regex_detector = RegexDetector()
        self.ner_detector = NERDetector(model_name=spacy_model)
        self.address_detector = AddressDetector()

        # Initialize replacer
        self.replacer = ConsistentReplacer(seed=42)

        # Track all detections for reporting
        self.detection_log: List[Dict] = []

        # Known names from context (populated during first pass)
        self._known_names: set = set()

        print("[✓] Engine initialized successfully.\n")

    def redact_document(self, input_path: str, output_path: str) -> Dict:
        """
        Redact all PII from a .docx document and save the result.

        Args:
            input_path: Path to the input .docx file
            output_path: Path for the redacted output .docx file

        Returns:
            Dictionary with redaction statistics
        """
        print(f"[*] Reading document: {input_path}")
        doc = Document(input_path)

        # Phase 1: Extract all text for context-aware detection
        full_text = self._extract_full_text(doc)
        print(f"[*] Extracted {len(full_text)} characters of text")

        # Phase 2: Detect PII in the full text to build the replacement map
        print("[*] Running PII detection pipeline...")
        entities = self._detect_all_pii(full_text)
        print(f"[✓] Detected {len(entities)} PII entities")

        # Phase 3: Build the replacement mapping
        replacement_map = self._build_replacement_map(entities)
        print(f"[✓] Built {len(replacement_map)} unique replacements")

        # Phase 4: Apply replacements to the document
        print("[*] Applying redactions to document...")
        stats = self._apply_redactions(doc, replacement_map)

        # Phase 5: Save the redacted document
        doc.save(output_path)
        print(f"[✓] Redacted document saved: {output_path}")

        return stats

    def _extract_full_text(self, doc: Document) -> str:
        """Extract all text from the document (paragraphs + tables)."""
        texts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)

        return '\n'.join(texts)

    def _detect_all_pii(self, text: str) -> List[PIIEntity]:
        """
        Run all detectors on the text and merge results.

        Detection order matters for overlap resolution:
        1. Regex (highest confidence for structured data)
        2. NER (good for names/organizations)
        3. Address (contextual, depends on keyword + location combos)
        """
        all_entities = []

        # Run regex detector
        regex_entities = self.regex_detector.detect(text)
        all_entities.extend(regex_entities)
        print(f"  [Regex]   Found {len(regex_entities)} entities")

        # Run NER detector
        ner_entities = self.ner_detector.detect(text)
        all_entities.extend(ner_entities)
        print(f"  [NER]     Found {len(ner_entities)} entities")

        # Run address detector
        addr_entities = self.address_detector.detect(text)
        all_entities.extend(addr_entities)
        print(f"  [Address] Found {len(addr_entities)} entities")

        # Context enhancement: find names that appear multiple times
        self._enhance_with_context(text, all_entities)

        # Merge overlapping detections
        merged = merge_overlapping_entities(all_entities)
        print(f"  [Merged]  {len(merged)} unique entities after deduplication")

        # Log all detections
        for entity in merged:
            self.detection_log.append({
                "text": entity.text,
                "type": entity.pii_type,
                "confidence": entity.confidence,
                "source": entity.source,
                "start": entity.start,
                "end": entity.end,
            })

        return merged

    def _enhance_with_context(self, text: str, entities: List[PIIEntity]):
        """
        Use contextual clues to improve detection:
        - Names mentioned near "Submitted By", "Name:", etc.
        - Enrollment numbers near "Enroll", "Roll No"
        - Additional names found through document structure
        """
        # Collect known person names from NER
        ner_names = {e.text for e in entities if e.pii_type == "NAME"}

        # Look for names after common prefixes
        # Use [^\S\n]+ instead of \s+ to prevent matching across newlines
        name_context_patterns = [
            re.compile(r'(?:Submitted\s+(?:By|To)|Supervisor|Guide|Mentor|Director|Professor|Mr\.|Mrs\.|Ms\.|Dr\.)\s*[:\s]*([A-Z][a-z]+(?:[^\S\n]+[A-Z][a-z]+)+)', re.MULTILINE),
            re.compile(r'(?:ownerName|employee_signature|name)\s*["\s:]+\s*["\']?([A-Z][a-z]+(?:[^\S\n]+[A-Z][a-z]+)+)', re.MULTILINE),
        ]

        # Known non-name titles/roles that might match the pattern
        skip_names = {
            "assistant professor", "technical director", "computer science",
            "bachelor of technology", "department of computer",
        }

        for pattern in name_context_patterns:
            for match in pattern.finditer(text):
                name = match.group(1).strip()
                if (name and name not in ner_names and len(name) > 3
                        and name.lower() not in skip_names):
                    entities.append(PIIEntity(
                        start=match.start(1),
                        end=match.end(1),
                        text=name,
                        pii_type="NAME",
                        confidence=0.85,
                        source="context"
                    ))
                    self._known_names.add(name)

    def _build_replacement_map(self, entities: List[PIIEntity]) -> Dict[str, Tuple[str, str]]:
        """
        Build a mapping from original PII text to (fake_replacement, pii_type).
        This ensures consistent replacement throughout the document.
        """
        replacement_map: Dict[str, Tuple[str, str]] = {}

        # Sort entities: process names first (so emails can reference fake names)
        sorted_entities = sorted(entities, key=lambda e: (
            0 if e.pii_type == "NAME" else
            1 if e.pii_type == "ORGANIZATION" else
            2 if e.pii_type == "EMAIL" else 3
        ))

        for entity in sorted_entities:
            key = entity.text.strip()
            if key and key not in replacement_map:
                fake = self.replacer.replace(key, entity.pii_type)
                replacement_map[key] = (fake, entity.pii_type)

        return replacement_map

    def _apply_redactions(self, doc: Document, replacement_map: Dict[str, Tuple[str, str]]) -> Dict:
        """
        Apply the replacement map to all text in the document.
        Handles paragraphs, tables, headers, and footers.

        Returns redaction statistics.
        """
        stats = {
            "total_replacements": 0,
            "by_type": {},
            "paragraphs_modified": 0,
            "tables_modified": 0,
        }

        # Sort replacements by length (longest first) to avoid partial matches
        sorted_replacements = sorted(
            replacement_map.items(),
            key=lambda x: len(x[0]),
            reverse=True
        )

        # Process paragraphs
        for para in doc.paragraphs:
            modified = self._replace_in_paragraph(para, sorted_replacements, stats)
            if modified:
                stats["paragraphs_modified"] += 1

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        modified = self._replace_in_paragraph(para, sorted_replacements, stats)
                        if modified:
                            stats["tables_modified"] += 1

        # Process headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    self._replace_in_paragraph(para, sorted_replacements, stats)
            if section.footer:
                for para in section.footer.paragraphs:
                    self._replace_in_paragraph(para, sorted_replacements, stats)

        return stats

    def _replace_in_paragraph(self, para, sorted_replacements, stats) -> bool:
        """
        Replace PII in a paragraph while preserving formatting.

        Strategy: We operate on runs (formatting units) to preserve bold, italic, etc.
        For each run, we check if any PII text appears and replace it.
        For PII that spans multiple runs, we concatenate runs and re-split.
        """
        modified = False
        full_text = para.text

        if not full_text.strip():
            return False

        # Check if any replacement is needed in this paragraph
        needs_replacement = False
        for original, (fake, pii_type) in sorted_replacements:
            if original in full_text:
                needs_replacement = True
                break

        if not needs_replacement:
            return False

        # Apply replacements to the full paragraph text
        new_text = full_text
        for original, (fake, pii_type) in sorted_replacements:
            count = new_text.count(original)
            if count > 0:
                new_text = new_text.replace(original, fake)
                stats["total_replacements"] += count
                stats["by_type"][pii_type] = stats["by_type"].get(pii_type, 0) + count
                modified = True

        # If text changed, update the paragraph runs
        if modified and new_text != full_text:
            self._update_paragraph_text(para, new_text)

        return modified

    def _update_paragraph_text(self, para, new_text: str):
        """
        Update paragraph text while trying to preserve formatting.

        Approach: If the paragraph has a single run, just update its text.
        If it has multiple runs, we update the first run with the full new text
        and clear the remaining runs. This may lose some inline formatting
        for cross-run PII, but ensures correctness.
        """
        if not para.runs:
            return

        if len(para.runs) == 1:
            para.runs[0].text = new_text
        else:
            # Preserve the formatting of the first run
            # Apply the entire new text to the first run
            # Clear all subsequent runs
            # This is the safest approach for maintaining document structure

            # Alternative: try to do run-by-run replacement
            old_text = para.text
            if old_text == new_text:
                return

            # Build a character-level mapping from old text to runs
            # Then try to apply changes preserving run boundaries
            run_texts = [run.text for run in para.runs]
            accumulated = ''
            run_boundaries = []  # (start, end) in full text for each run
            for rt in run_texts:
                start = len(accumulated)
                accumulated += rt
                run_boundaries.append((start, len(accumulated)))

            # Simple approach: set first run to new text, clear others
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''

    def get_detection_summary(self) -> Dict:
        """Get a summary of all detections for evaluation."""
        summary = {
            "total_entities": len(self.detection_log),
            "by_type": {},
            "by_source": {},
            "entities": self.detection_log,
        }

        for entry in self.detection_log:
            ptype = entry["type"]
            source = entry["source"]
            summary["by_type"][ptype] = summary["by_type"].get(ptype, 0) + 1
            summary["by_source"][source] = summary["by_source"].get(source, 0) + 1

        return summary


def print_summary(stats: Dict, detection_summary: Dict):
    """Print a formatted summary of the redaction process."""
    print("\n" + "=" * 60)
    print("  PII REDACTION SUMMARY")
    print("=" * 60)

    print(f"\n  Total PII entities detected: {detection_summary['total_entities']}")
    print(f"  Total replacements made:     {stats['total_replacements']}")
    print(f"  Paragraphs modified:         {stats['paragraphs_modified']}")
    print(f"  Table cells modified:        {stats['tables_modified']}")

    print("\n  Detection by PII Type:")
    print("  " + "-" * 40)
    for ptype, count in sorted(detection_summary["by_type"].items()):
        print(f"    {ptype:<25} {count:>5}")

    print("\n  Detection by Source:")
    print("  " + "-" * 40)
    for source, count in sorted(detection_summary["by_source"].items()):
        print(f"    {source:<25} {count:>5}")

    if stats["by_type"]:
        print("\n  Replacements by PII Type:")
        print("  " + "-" * 40)
        for ptype, count in sorted(stats["by_type"].items()):
            print(f"    {ptype:<25} {count:>5}")

    print("\n" + "=" * 60)


def save_detection_report(detection_summary: Dict, output_path: str):
    """Save detailed detection report as JSON."""
    report_path = output_path.replace('.docx', '_detection_report.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(detection_summary, f, indent=2, ensure_ascii=False)
    print(f"[✓] Detection report saved: {report_path}")


def main():
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool - Detects and replaces personally identifiable information in .docx files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python pii_redactor.py report.docx
    python pii_redactor.py report.docx --output redacted_report.docx
    python pii_redactor.py report.docx --output redacted.docx --report
        """
    )

    parser.add_argument(
        "input",
        help="Path to the input .docx file"
    )
    parser.add_argument(
        "--output", "-o",
        help="Path for the redacted output .docx file (default: <input>_redacted.docx)",
        default=None
    )
    parser.add_argument(
        "--report", "-r",
        help="Save a detailed detection report as JSON",
        action="store_true"
    )
    parser.add_argument(
        "--model", "-m",
        help="spaCy model name (default: en_core_web_sm)",
        default="en_core_web_sm"
    )

    args = parser.parse_args()

    # Validate input file
    if not os.path.isfile(args.input):
        print(f"[ERROR] Input file not found: {args.input}")
        sys.exit(1)

    if not args.input.lower().endswith('.docx'):
        print("[ERROR] Input file must be a .docx file")
        sys.exit(1)

    # Set default output path
    if args.output is None:
        base, ext = os.path.splitext(args.input)
        args.output = f"{base}_redacted{ext}"

    # Run redaction
    start_time = time.time()

    redactor = PIIRedactor(spacy_model=args.model)
    stats = redactor.redact_document(args.input, args.output)
    detection_summary = redactor.get_detection_summary()

    elapsed = time.time() - start_time

    # Print summary
    print_summary(stats, detection_summary)
    print(f"\n  Time elapsed: {elapsed:.2f} seconds")

    # Save report if requested
    if args.report:
        save_detection_report(detection_summary, args.output)

    # Print sample replacements
    if detection_summary["entities"]:
        print("\n  Sample Replacements:")
        print("  " + "-" * 55)
        seen_types = set()
        for entry in detection_summary["entities"]:
            if entry["type"] not in seen_types and len(seen_types) < 10:
                original = entry["text"]
                fake = redactor.replacer._cache.get(
                    f"{entry['type']}:{original.lower().strip()}", "[REDACTED]"
                )
                print(f"    {entry['type']:<20} {original[:30]:<30} → {fake}")
                seen_types.add(entry["type"])

    print()


if __name__ == "__main__":
    main()
